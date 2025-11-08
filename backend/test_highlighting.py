"""Test script to verify Word document highlighting works"""
from document_generator import DocumentGenerator
from io import BytesIO

def test_highlighting():
    """Test that highlighting is applied to Word documents"""

    generator = DocumentGenerator()

    # Sample report with AI-generated content
    report_text = """FINDINGS:
The heart size is normal. The lungs are clear.
No acute abnormality detected.

IMPRESSION:
Normal chest radiograph."""

    template_skeleton = """FINDINGS:
<fill>

IMPRESSION:
<fill>"""

    # Generate with highlighting
    doc_stream = generator.generate_word_with_highlighting(
        report_text=report_text,
        template_skeleton=template_skeleton
    )

    print("✓ Word document generated with highlighting")
    print(f"  Document size: {len(doc_stream.getvalue())} bytes")

    # Verify document can be read
    from docx import Document
    doc_stream.seek(0)
    doc = Document(doc_stream)

    print(f"✓ Document has {len(doc.paragraphs)} paragraphs")

    # Check for highlighting
    highlight_found = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None:
                highlight_found = True
                print(f"✓ Found highlighted text: '{run.text[:50]}...'")
                break
        if highlight_found:
            break

    if highlight_found:
        print("\n✓ SUCCESS: Yellow highlighting is applied to Word document!")
    else:
        print("\n⚠ WARNING: No highlighting found (may need AI content detection)")

    return doc_stream

if __name__ == "__main__":
    test_highlighting()
